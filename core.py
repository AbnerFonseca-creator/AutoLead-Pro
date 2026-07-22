# core.py
import os
from dotenv import load_dotenv
import google.generativeai as genai
import pandas as pd
from docx import Document
import geocoder
import requests

# Carrega as variáveis ocultas do arquivo .env (Blindagem de Segurança)
load_dotenv()

# ==============================================================================
# CONFIGURAÇÕES (APENAS GEMINI, SEM GOOGLE MAPS)
# ==============================================================================
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    modelo_ia = genai.GenerativeModel('gemini-2.5-flash')
else:
    print("ALERTA: Chave do Gemini não encontrada no arquivo .env!")

# ==============================================================================
# MOTOR DE BUSCA GRATUITO (OPENSTREETMAP / OVERPASS API)
# ==============================================================================
def obter_localizacao_atual():
    """Rastreia a localização atual via IP com sistema de segurança (Fallback)."""
    try:
        g = geocoder.ip('me')
        if g.latlng:
            return g.latlng[0], g.latlng[1], g.city
    except:
        pass
    
    # Se o firewall do PC bloquear o rastreio, usa coordenadas seguras de base
    return -19.9208, -43.9378, "Belo Horizonte"

def buscar_empresas_google(termo, lat, lng, raio_km):
    """Varre o OpenStreetMap trazendo dados e coordenadas para os Pinos do Mapa."""
    raio_metros = int(raio_km * 1000)
    
    # Adicionado "amenity" à query para encontrar mais tipos de comércios
    query = f"""
    [out:json][timeout:25];
    (
      nwr["name"~"{termo}",i](around:{raio_metros},{lat},{lng});
      nwr["shop"~"{termo}",i](around:{raio_metros},{lat},{lng});
      nwr["amenity"~"{termo}",i](around:{raio_metros},{lat},{lng});
    );
    out center;
    """
    
    url = "http://overpass-api.de/api/interpreter"
    headers = {'User-Agent': 'AutoLeadPro_App/2.0'}
    
    try:
        resposta = requests.post(url, data={'data': query}, headers=headers)
        
        if resposta.status_code != 200:
            print(f"Erro na API Gratuita: Servidor sobrecarregado (HTTP {resposta.status_code})")
            return []
            
        dados_json = resposta.json()
        resultados = dados_json.get('elements', [])
        
        empresas = []
        
        # Aumentei o limite para 15 resultados para popular o mapa
        for local in resultados[:15]:
            tags = local.get('tags', {})
            nome = tags.get('name', 'Empresa Sem Nome')
            
            # Pula empresas que não têm nome registrado no mapa
            if nome == 'Empresa Sem Nome':
                continue
                
            rua = tags.get('addr:street', '')
            num = tags.get('addr:housenumber', '')
            endereco = f"{rua}, {num}".strip(", ")
            if not endereco:
                endereco = "Endereço não detalhado no mapa"
                
            telefone = tags.get('phone', tags.get('contact:phone', 'Telefone não público'))
            
            # Lógica NOVA: Captura lat/lon para colocar os alfinetes no TkinterMapView
            lat_empresa = local.get('lat', local.get('center', {}).get('lat', lat))
            lon_empresa = local.get('lon', local.get('center', {}).get('lon', lng))
            
            empresas.append({
                "Nome": nome,
                "Endereço": endereco,
                "Nota": "N/A (OSM)",
                "Telefone": telefone,
                "lat": lat_empresa,
                "lon": lon_empresa
            })
            
        return empresas
    except Exception as e:
        print(f"Falha de conexão com o mapa: {e}")
        return []

# ==============================================================================
# MOTOR DE INTELIGÊNCIA ARTIFICIAL E EXPORTAÇÃO
# ==============================================================================
def gerar_mensagem(nome_alvo, nome_empresa_usuario, api_key_usuario):
    """Cria a abordagem de vendas via IA usando a chave fornecida na interface."""
    if not api_key_usuario:
        return "Erro: Chave API do Gemini não informada no painel."
        
    try:
        # Configura a IA no momento do disparo com a chave do usuário
        genai.configure(api_key=api_key_usuario)
        modelo_ia = genai.GenerativeModel('gemini-2.5-flash')
        
        prompt = f"""
        Aja como um representante comercial e anunciante da empresa '{nome_empresa_usuario}'.
        Escreva uma mensagem curta, profissional e direta (máximo 2 parágrafos) para o WhatsApp da empresa '{nome_alvo}'.
        
        REGRAS ESTRITAS: 
        1. PROIBIDO usar colchetes como [Seu Nome] ou [Nome do Gerente]. 
        2. Não invente nomes próprios de pessoas.
        3. Inicie cumprimentando a empresa '{nome_alvo}' diretamente.
        4. O objetivo é apresentar os serviços/produtos da '{nome_empresa_usuario}' e iniciar uma negociação.
        """
        return modelo_ia.generate_content(prompt).text
    except Exception as e:
        return f"Erro ao gerar mensagem com a IA: Verifique se sua chave é válida."

def exportar_excel(dados):
    # Retira as colunas de latitude/longitude antes de gerar o Excel para não sujar a planilha do usuário
    dados_limpos = [{k: v for k, v in d.items() if k not in ['lat', 'lon']} for d in dados]
    df = pd.DataFrame(dados_limpos)
    df.to_excel("leads_geolocalizados.xlsx", index=False)

def exportar_word(dados):
    doc = Document()
    doc.add_heading('Leads Geolocalizados', 0)
    for emp in dados:
        doc.add_heading(emp['Nome'], level=1)
        doc.add_paragraph(f"Endereço: {emp['Endereço']} | Nota: {emp.get('Nota', '')}")
        doc.add_paragraph("Mensagem Sugerida:")
        doc.add_paragraph(emp.get('Mensagem Pronta', ''))
    doc.save("leads_geolocalizados.docx")

def disparar_webhook(webhook_url, dados_lead):
    if not webhook_url:
        return False, "URL não configurada."
    
    try:
        # Monta um pacote JSON padronizado para o robô ler
        payload = {
            "nome": dados_lead.get("Nome", ""),
            "telefone": dados_lead.get("Telefone", ""),
            "endereco": dados_lead.get("Endereço", ""),
            "mensagem": dados_lead.get("Mensagem Pronta", "")
        }
        
        headers = {'Content-Type': 'application/json'}
        resposta = requests.post(webhook_url, json=payload, headers=headers, timeout=10)
        
        # Códigos 200 ou 201 significam que o robô recebeu com sucesso
        if resposta.status_code in [200, 201]:
            return True, "Enviado com sucesso!"
        else:
            return False, f"Erro HTTP {resposta.status_code}"
            
    except Exception as e:
        return False, f"Falha de conexão: {e}"